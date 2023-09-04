import dicts
import os
import docx
import openpyxl


def clear():
    os.chdir('C:\\Users\\Ryzhk\\Desktop\\Грузоперевозки')
    try:
        for i in os.listdir():
            os.remove(i)
    except PermissionError:
        print(f"Файл {i} открыт, закрой его")


class Docx:
    def __init__(self, way_list: list, date: str, doc_num: int, organization: str):
        self.organization = organization
        self.doc_num = doc_num
        self.date = date
        self.street_1 = way_list[1]
        self.street_2 = way_list[2]
        self.money = way_list[3]
        self.tons = way_list[4]
        self.cargo = way_list[5]
        self.doc = docx.Document(
            'C:\\Users\\Ryzhk\\PycharmProjects\\Грузоперевозки\\CargoTransportation\\Шаблон_dox.docx')

    def word_num(self):
        if self.money == "{Неверная сумма}":
            return "{Неверная сумма}"
        else:
            self.money = int(self.money)
            return f"{self.money},0 ({dicts.B12_thousands_dict[self.money // 1000]} {dicts.B12_thousands_two_dict[self.money // 1000]}{dicts.B12_hundreds_dict[(self.money // 100) % 10]}) руб 00коп"

    def docx_writer(self):
        for i in dicts.world_dates:
            self.doc.tables[i[0]].rows[i[1]].cells[i[2]].text = self.date
        self.doc.tables[0].rows[5].cells[1].text = dicts.organization_dict[self.organization]
        self.doc.tables[0].rows[5].cells[6].text = dicts.organization_dict[self.organization]
        self.doc.tables[0].rows[8].cells[1].text = self.cargo
        self.doc.tables[0].rows[12].cells[1].text = self.tons
        self.doc.tables[1].rows[0].cells[1].text = self.street_1
        self.doc.tables[1].rows[0].cells[6].text = self.street_2
        self.doc.tables[1].rows[8].cells[3].text = self.tons
        self.doc.tables[1].rows[8].cells[8].text = self.tons
        self.doc.tables[3].rows[28].cells[1].text = self.word_num()
        self.doc.tables[3].rows[37].cells[1].text = dicts.organization_short_dict[self.organization]

        self.doc.save(
            f'C:\\Users\\Ryzhk\\Desktop\\{self.organization} архив\\' + str(self.doc_num) + '      ' + self.date + '.docx')


class Exel:

    def __init__(self, list_with_str: list, act_val: int, document_number: int, organization: str, money_at_hour: int):
        self.date = list_with_str[0][0]
        self.list_with_str = list_with_str
        self.act_val = act_val
        self.document_number = document_number
        self.organization = organization
        self.wb = openpyxl.load_workbook(
            'C:\\Users\\Ryzhk\\PycharmProjects\\Грузоперевозки\\CargoTransportation\\Шаблон_exel.xlsx')
        self.ws = self.wb.active
        self.money_at_hour = money_at_hour
        self.check()

    def check(self):
        c = 0
        # if len(self.list_with_str) != 5:
        # self.list_with_str = ['{Неверная дата}', "Неверная улица}", "Неверная улица}", "{Неверная сумма}", "{Неверные тонны}", "{Неверный груз}"]
        for i in self.list_with_str:
            if not ((c != 0 and i[0] == "") or (c == 0 and (len(i[0]) == 10 and i[0][0:2].isdigit() and i[0][3:5].isdigit() and i[0][6:].isdigit() and i[0][2] == "." and i[0][5] == "."))):
                i[0] = "0" + i[0]
                if c == 0 and (len(i[0]) == 10 and i[0][0:2].isdigit() and i[0][3:5].isdigit() and i[0][6:].isdigit() and i[0][2] == "." and i[0][5] == "."):
                    self.date = "0" + self.date
                else:
                    print("\033[31m ", self.act_val, f"Неверная дата {self.date}")
                    self.date = "{Неверная дата}"
                    i[0] = "{Неверная дата}"
            if len(i[1]) == 0:
                print("\033[31m ", self.act_val, f"Неверная улица {i[1]}")
                i[1] = "{Неверная улица}"
            if len(i[2]) == 0:
                print("\033[31m ", self.act_val, f"Неверная улица {i[1]}")
                i[2] = "{Неверная улица}"
            if len(i[5]) == 0:
                print("\033[31m ", self.act_val, f"Неверный груз {i[5]}")
                i[5] = "{Неверный груз}"
            if not (i[3].isdigit() and 1000 <= int(i[3]) < 110000 and int(i[3]) % 100 == 0):
                print("\033[31m ", self.act_val, f"Неверная сумма {i[3]}")
                i[3] = "{Неверная сумма}"
            if not (i[4].isdigit() and 1 <= int(i[4]) <= 10):
                print("\033[31m ", self.act_val, f"Неверные тонны {i[4]}")
                i[4] = "{Неверные тонны}"
            c += 1

    def writer(self):
        street_string = ''
        money = 0
        for day_list in self.list_with_str:
            street_string += day_list[1] + ' - ' + day_list[2] + ';\n'
            if day_list[3] != "{Неверная сумма}":
                money += int(day_list[3])
        return money, street_string

    def b12(self, number):
        if 1000 <= number < 110000:
            return f"{dicts.B12_thousands_dict[number // 1000]} {dicts.B12_thousands_two_dict[number // 1000]}{dicts.B12_hundreds_dict[(number // 100) % 10]} рублей 00коп "
        else:
            return "Неверная сумма"

    def exel_writer(self):
        money, street_string = self.writer()
        if self.date == "{Неверная дата}":
            self.ws['B3'] = "Акт № " + str(self.act_val) + " от " + "Неверной даты"
            self.ws['D12'] = 'Неверная дата Услуги манипулятора (МАРКА А/М MAN, Х 012 ХВ,' + \
                             'Шестернин Александр Олегович) по маршруту: \n' + street_string
        else:
            self.ws['B3'] = "Акт № " + str(self.act_val) + " от " + self.date[:2] + " " \
                            + dicts.month_dict[int(self.date[3:5])] + " " + self.date[6:] + " " + "г."
            self.ws['D12'] = self.date + 'г. Услуги манипулятора (МАРКА А/М MAN, Х 012 ХВ,' + \
                             'Шестернин Александр Олегович) по маршруту: \n' + street_string
        if money % self.money_at_hour == 0:
            self.ws['U12'] = str(int(money / self.money_at_hour))
            self.ws['Z12'] = str(self.money_at_hour)
        else:
            self.ws['U12'] = "{Неверные часы}"
            print("\033[31m ", self.act_val, "Сумма не делится на количество денег зарабатываемых в час")

        self.ws['AD12'] = str(money) + ',00'
        self.ws['AD14'] = str(money) + ',00'
        self.ws['B17'] = f"Всего оказано услуг 1 на сумму {money},00"
        self.ws['B18'] = self.b12(money)
        self.ws['F7'] = dicts.organization_dict[self.organization]
        os.chdir(f'C:\\Users\\Ryzhk\\Desktop\\{self.organization} архив')
        self.wb.save(str(self.document_number) + '      ' + self.date + ' Акт №' + str(self.act_val) + '.xlsx')
